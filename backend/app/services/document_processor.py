# app/services/document_processor.py

from typing import Tuple, Dict, Any, BinaryIO, NamedTuple
import os
import aiofiles
import hashlib
from datetime import datetime
from uuid import UUID, uuid4
from fastapi import UploadFile, File

from app.models.documents.source_document import SourceDocument

# Conditional imports to handle missing dependencies gracefully
try:
    import fitz  # PyMuPDF for PDF handling
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False
    print("WARNING: PyMuPDF (fitz) not installed. PDF processing will be limited.")

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("WARNING: python-docx not installed. DOCX processing will be limited.")

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False
    print("WARNING: chardet not installed. Text encoding detection will be limited.")


class ProcessResult(NamedTuple):
    """文档处理结果"""
    document: SourceDocument
    text_content: str
    metadata: Dict[str, Any]
    error: str = None


class DocumentProcessor:
    """增强版文档处理器，支持多格式文档解析与结构保留"""
    
    def __init__(self, documents_dir: str = "./data/documents", archives_dir: str = "./data/archives"):
        """初始化文档处理器
        
        Args:
            documents_dir: 文档存储目录
            archives_dir: 归档存储目录
        """
        self.documents_dir = documents_dir
        self.archives_dir = archives_dir
        
        # 确保目录存在
        os.makedirs(documents_dir, exist_ok=True)
        os.makedirs(archives_dir, exist_ok=True)
    
    async def process_file(self, file: BinaryIO, filename: str) -> ProcessResult:
        """处理上传的文件
        
        Args:
            file: 文件对象
            filename: 文件名
            
        Returns:
            包含处理结果的ProcessResult对象
        """
        try:
            # 打印日志以进行调试
            print(f"Processing file: {filename}")
            
            # 获取文件类型
            file_ext = os.path.splitext(filename)[1].lower().lstrip('.')
            if not file_ext:
                print(f"Missing file extension: {filename}")
                return ProcessResult(
                    document=None, 
                    text_content="", 
                    metadata={}, 
                    error=f"Missing file extension in filename: {filename}"
                )
            
            # 生成唯一文件名
            unique_id = str(uuid4())
            safe_filename = f"{unique_id}_{os.path.basename(filename)}"
            file_path = os.path.join(self.documents_dir, safe_filename)
            
            # 确保目录存在
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # 保存文件 - 更安全的方法读取文件内容
            content = None
            try:
                # 判断文件对象类型
                if hasattr(file, "read"):
                    # 对应于常规文件对象，使用标准读取
                    content = await file.read()
                else:
                    print(f"Warning: file object doesn't have 'read' method, type: {type(file)}")
                    # 尝试直接使用文件对象，如果它本身就是字节内容
                    content = file
                    
                if not content:
                    raise ValueError("Could not read file content")
                
                # 确保内容是字节类型
                if not isinstance(content, bytes):
                    if isinstance(content, str):
                        content = content.encode('utf-8')
                    else:
                        raise TypeError(f"Unexpected content type: {type(content)}")
            except Exception as e:
                print(f"Error reading file content: {str(e)}")
                return ProcessResult(
                    document=None, 
                    text_content="", 
                    metadata={}, 
                    error=f"Error reading file: {str(e)}"
                )
            
            # 计算文件内容哈希
            content_hash = f"sha256:{hashlib.sha256(content).hexdigest()}"
            
            # 写入文件
            try:
                async with aiofiles.open(file_path, 'wb') as f:
                    await f.write(content)
                print(f"File written to {file_path}")
            except Exception as e:
                print(f"Error writing file to disk: {str(e)}")
                return ProcessResult(
                    document=None, 
                    text_content="", 
                    metadata={}, 
                    error=f"Error writing file to disk: {str(e)}"
                )
            
            # 重置文件指针以便后续读取 - 我们已经有了内容，所以不再需要此步骤
            # 创建SourceDocument记录
            document = SourceDocument(
                id=UUID(unique_id),
                title=filename,
                type=file_ext,
                content_hash=content_hash,
                file_path=file_path,
                metadata={},
                accessed_at=datetime.now()
            )
            
            # 提取内容和元数据
            text_content = ""
            metadata = {}
            try:
                text_content, metadata = await self._extract_content_and_metadata(file_path, file_ext)
                print(f"Extracted content length: {len(text_content)}, metadata keys: {list(metadata.keys())}")
            except Exception as e:
                print(f"Error in content extraction: {str(e)}")
                # 简单地使用文本编码作为备选方案
                try:
                    # 尝试将内容解码为文本
                    text_content = content.decode('utf-8', errors='replace')
                    metadata = {
                        'file_size': len(content),
                        'note': 'Extracted as plain text due to processing error'
                    }
                    print(f"Fallback extraction: using content as plain text, size: {len(text_content)}")
                except Exception as decode_err:
                    print(f"Error in fallback content extraction: {str(decode_err)}")
                    return ProcessResult(
                        document=None, 
                        text_content="", 
                        metadata={}, 
                        error=f"Error extracting content and metadata: {str(e)}. Fallback also failed: {str(decode_err)}"
                    )
            
            # 检查元数据中是否有错误信息
            if isinstance(metadata, dict) and "error" in metadata:
                print(f"Error reported in metadata: {metadata['error']}")
                # 使用文本编码作为备选方案，而不是直接返回错误
                try:
                    # 尝试将内容解码为文本
                    text_content = content.decode('utf-8', errors='replace')
                    metadata = {
                        'file_size': len(content),
                        'note': f'Extracted as plain text due to processing error: {metadata["error"]}',
                        'original_error': metadata["error"]
                    }
                    print(f"Using fallback extraction due to metadata error")
                except Exception as decode_err:
                    print(f"Fallback extraction failed: {str(decode_err)}")
                    return ProcessResult(
                        document=None, 
                        text_content="", 
                        metadata={}, 
                        error=metadata["error"]
                    )
            
            # 更新文档元数据
            document.metadata.update(metadata)
            
            # 创建归档副本
            archive_path = os.path.join(self.archives_dir, f"{unique_id}.{file_ext}")
            
            # 确保归档目录存在
            os.makedirs(os.path.dirname(archive_path), exist_ok=True)
            
            try:
                await self._create_archive_copy(file_path, archive_path)
                document.archived_path = archive_path
                print(f"Created archive copy at {archive_path}")
            except Exception as e:
                # 如果创建归档副本失败，不要中断处理
                print(f"Warning: Could not create archive copy: {e}")
                # 仍然继续处理
            
            print(f"File processing completed successfully for {filename}")
            return ProcessResult(document=document, text_content=text_content, metadata=metadata)
            
        except Exception as e:
            # 出错时返回错误信息
            import traceback
            traceback.print_exc() # 这将打印详细的堆栈跟踪到控制台
            print(f"Unhandled error in document processing: {str(e)}")
            return ProcessResult(
                document=None, 
                text_content="", 
                metadata={}, 
                error=f"Error processing file: {str(e)}"
            )
    
    async def process_url(self, url: str) -> ProcessResult:
        """处理URL
        
        Args:
            url: 要处理的URL
            
        Returns:
            包含处理结果的ProcessResult对象
        """
        try:
            from app.services.web_archiver import WebArchiver
            
            # 创建网页归档器
            archiver = WebArchiver(self.archives_dir)
            
            # 捕获和归档网页
            result = await archiver.capture_and_archive(url)
            
            # 创建SourceDocument记录
            document = SourceDocument(
                id=UUID(result["archive_id"]),
                title=result.get("metadata", {}).get("title", url),
                type="webpage",
                content_hash=result["content_hash"],
                url=url,
                archived_path=result["warc_path"],
                metadata=result["metadata"],
                accessed_at=datetime.now()
            )
            
            return ProcessResult(
                document=document,
                text_content=result["content"],
                metadata=result["metadata"]
            )
            
        except Exception as e:
            # 出错时返回错误信息
            import traceback
            traceback.print_exc()
            return ProcessResult(
                document=None, 
                text_content="", 
                metadata={}, 
                error=f"Error processing URL: {str(e)}"
            )
    
    async def _create_archive_copy(self, source_path: str, target_path: str) -> bool:
        """创建文件的归档副本"""
        try:
            # 简单复制文件作为归档
            async with aiofiles.open(source_path, 'rb') as src_file:
                content = await src_file.read()
                
            async with aiofiles.open(target_path, 'wb') as dst_file:
                await dst_file.write(content)
                
            return True
        except Exception as e:
            print(f"Error creating archive copy: {e}")
            return False
    
    async def _extract_content_and_metadata(self, file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """提取文件内容和元数据"""
        if file_type == 'pdf':
            return await self._process_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return await self._process_word(file_path)
        elif file_type == 'txt':
            return await self._process_text(file_path)
        else:
            # 对于未知类型，尝试作为文本处理
            try:
                return await self._process_text(file_path)
            except Exception as e:
                return "", {"error": f"Unsupported file type: {file_type}. Error: {str(e)}"}
    
    async def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """处理PDF文件，提取文本和结构信息"""
        if not HAS_FITZ:
            return "", {"error": "PyMuPDF (fitz) library not installed. Unable to process PDF files."}
        
        try:
            doc = fitz.open(file_path)
            
            # 提取元数据
            metadata = {
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'keywords': doc.metadata.get('keywords', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'creation_date': doc.metadata.get('creationDate', ''),
                'modification_date': doc.metadata.get('modDate', ''),
                'page_count': len(doc),
                'file_size': os.path.getsize(file_path),
            }
            
            # 提取结构化文本
            text_blocks = []
            try:
                toc = doc.get_toc()  # 获取目录结构
            except Exception:
                toc = []  # 如果获取目录失败，使用空列表
            
            # 处理每一页，保留段落和布局信息
            full_text = ""
            for page_num, page in enumerate(doc):
                try:
                    page_text = page.get_text()
                    full_text += page_text + "\n\n"
                    
                    # 获取更详细的块信息用于元数据
                    try:
                        blocks = page.get_text("dict")["blocks"]
                        for block in blocks:
                            if block["type"] == 0:  # 文本块
                                for line in block["lines"]:
                                    line_text = " ".join(span["text"] for span in line["spans"])
                                    font_info = line["spans"][0] if line["spans"] else {}
                                    
                                    # 记录字体、大小等排版信息，用于判断标题或正文
                                    text_block_info = {
                                        "text": line_text,
                                        "page": page_num + 1,
                                        "font": font_info.get("font", ""),
                                        "size": font_info.get("size", 0),
                                        "position": (block["bbox"][0], block["bbox"][1])
                                    }
                                    
                                    text_blocks.append(text_block_info)
                    except Exception as e:
                        print(f"Warning: Error extracting detailed blocks from page {page_num}: {e}")
                        # 继续处理下一页
                except Exception as e:
                    print(f"Warning: Error processing page {page_num}: {e}")
                    # 继续处理下一页
            
            # 将结构信息添加到元数据
            metadata['structure'] = {
                'toc': toc,
                'text_blocks': text_blocks[:100],  # 只保存部分块以避免元数据过大
            }
            
            return full_text, metadata
        except Exception as e:
            # 返回错误信息
            return "", {"error": f"Error processing PDF: {str(e)}"}
    
    async def _process_word(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """处理Word文档，提取文本和结构信息"""
        if not HAS_DOCX:
            return "", {"error": "python-docx library not installed. Unable to process Word documents."}
        
        try:
            doc = docx.Document(file_path)
            
            # 提取元数据
            core_properties = doc.core_properties
            metadata = {
                'title': core_properties.title or '',
                'author': core_properties.author or '',
                'subject': core_properties.subject or '',
                'keywords': core_properties.keywords or '',
                'created': str(core_properties.created) if core_properties.created else '',
                'modified': str(core_properties.modified) if core_properties.modified else '',
                'paragraph_count': len(doc.paragraphs),
                'file_size': os.path.getsize(file_path),
            }
            
            # 提取结构化文本
            text_blocks = []
            full_text = ""
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    # 获取段落样式信息
                    style_name = para.style.name if para.style else "Normal"
                    
                    # 记录段落信息，包括样式、层级等
                    text_info = {
                        "text": para.text,
                        "index": i,
                        "style": style_name,
                        "level": 0 if "Heading" not in style_name else 
                                int(style_name.replace("Heading ", "")) 
                                if style_name.replace("Heading ", "").isdigit() else 0,
                        "char_offset": len(full_text)
                    }
                    
                    text_blocks.append(text_info)
                    full_text += para.text + "\n"
            
            # 将结构信息添加到元数据
            metadata['structure'] = {
                'paragraphs': text_blocks,
            }
            
            return full_text, metadata
        except Exception as e:
            # 返回错误信息
            return "", {"error": f"Error processing Word document: {str(e)}"}
    
    async def _process_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """处理纯文本文件"""
        try:
            # 读取文件内容
            async with aiofiles.open(file_path, 'rb') as f:
                content = await f.read()
            
            # 检测编码
            if HAS_CHARDET:
                encoding_result = chardet.detect(content)
                encoding = encoding_result['encoding'] or 'utf-8'
                encoding_confidence = encoding_result['confidence']
            else:
                encoding = 'utf-8'
                encoding_confidence = 1.0
            
            # 以检测到的编码解码文本
            try:
                text = content.decode(encoding)
            except UnicodeDecodeError:
                # 如果解码失败，尝试使用UTF-8
                text = content.decode('utf-8', errors='replace')
            
            # 创建元数据
            metadata = {
                'file_size': os.path.getsize(file_path),
                'encoding': encoding,
                'encoding_confidence': encoding_confidence,
                'line_count': text.count('\n') + 1,
                'char_count': len(text),
                'word_count': len(text.split())
            }
            
            return text, metadata
        except Exception as e:
            return "", {"error": f"Error processing text file: {str(e)}"}