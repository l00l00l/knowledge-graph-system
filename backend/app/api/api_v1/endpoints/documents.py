# 文件: app/api/api_v1/endpoints/documents.py
from typing import List, Optional
from uuid import UUID, uuid4
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query, BackgroundTasks
from app.models.documents.source_document import SourceDocument
from app.services.document_processor import DocumentProcessor
from app.db.neo4j_db import Neo4jDatabase
from app.api.deps import get_db
import os
import json
import aiofiles  # 确保在文件顶部导入这个库
from sqlalchemy.orm import Session
from app.db.sqlite_db import get_sqlite_db
from app.db.models import Document
from app.db.init_db import init_db
from fastapi.responses import FileResponse
from app.models.entities.entity import Entity

# 修正导入路径
from app.services.knowledge_extractor import SpacyNERExtractor

from pydantic import BaseModel
from datetime import datetime
from typing import Dict, Any, Optional

init_db()

class DocumentResponse(BaseModel):
    """文档响应模型"""
    id: str
    title: str
    type: str
    content_hash: str
    file_path: Optional[str] = None
    url: Optional[str] = None
    archived_path: Optional[str] = None
    metadata: Dict[str, Any] = {}
    accessed_at: datetime
    created_at: datetime
    updated_at: datetime

router = APIRouter()

# Create necessary directories
os.makedirs("./data/documents", exist_ok=True)
os.makedirs("./data/archives", exist_ok=True)

# Initialize mock documents for testing
mock_documents = []

# For development only: Mock implementation of knowledge extractor
class MockKnowledgeExtractor:
    """简化的知识抽取器用于测试"""
    
    def __init__(self, db):
        self.db = db
    
    async def extract_entities(self, document, text_content):
        print(f"Mock extracting entities from document: {document.title}")
        return []
    
    async def extract_relationships(self, document, entities, text_content):
        print(f"Mock extracting relationships from document: {document.title}")
        return []
    
    async def create_knowledge_traces(self, document, entities, relationships, text_content):
        print(f"Mock creating knowledge traces for document: {document.title}")
        return []

# Enhanced upload endpoint with better error handling

@router.post("/upload", response_model=Dict[str, Any])
async def upload_document(
    file: UploadFile = File(...),
    extract_knowledge: bool = Form(False),
    sqlite_db: Session = Depends(get_sqlite_db),
    neo4j_db: Neo4jDatabase = Depends(get_db)
):
    """上传文档并可选地提取知识"""
    try:
        # 创建文档处理器
        document_processor = DocumentProcessor()
        
        # 处理文档
        try:
            result = await document_processor.process_file(file.file, file.filename)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error processing document: {str(e)}")
        
        if result.error:
            raise HTTPException(status_code=400, detail=f"Error processing document: {result.error}")
        
        # 确保metadata是有效的JSON
        metadata_json = "{}"
        if result.document.metadata:
            try:
                # Make sure it's serializable
                metadata_json = json.dumps(result.document.metadata)
                # Verify it can be parsed back
                json.loads(metadata_json)
                print(f"Valid metadata JSON created: {metadata_json[:100]}...")
            except Exception as e:
                print(f"Error serializing metadata to JSON: {e}")
                metadata_json = "{}"
        
        # 将文档存入SQLite数据库
        doc_model = Document(
            id=str(result.document.id),
            title=result.document.title,
            type=result.document.type,
            content_hash=result.document.content_hash,
            file_path=result.document.file_path,
            url=result.document.url,
            archived_path=result.document.archived_path,
            doc_metadata=metadata_json  # 使用验证过的JSON
        )
        
        # 添加到数据库
        sqlite_db.add(doc_model)
        sqlite_db.commit()
        sqlite_db.refresh(doc_model)
        
        print(f"Document saved to SQLite with ID: {doc_model.id}")
        
        # 验证存储的JSON
        verification = sqlite_db.query(Document).filter(Document.id == doc_model.id).first()
        if verification:
            print(f"Successfully verified document in database: {verification.title}")
            print(f"Stored doc_metadata: {verification.doc_metadata[:100]}...")
        else:
            print(f"WARNING: Could not verify document {doc_model.id} in database")
        
        # 如果需要，提取知识
        entities = []
        relationships = []
        extract_error = None
        
        if extract_knowledge:
            try:
                # 创建知识抽取器 - 使用Neo4j
                extractor = SpacyNERExtractor(neo4j_db)
                
                # 提取实体
                entities = await extractor.extract_entities(result.document, result.text_content)
                
                # 提取关系
                if len(entities) >= 2:
                    relationships = await extractor.extract_relationships(result.document, entities, result.text_content)
                    
                    # 创建溯源记录
                    await extractor.create_knowledge_traces(result.document, entities, relationships, result.text_content)
                
            except Exception as e:
                extract_error = str(e)
                print(f"Error during knowledge extraction: {e}")
        # 在upload_document函数中添加创建基本实体的部分
        if not extract_knowledge:
            # 即使不提取知识，也创建一个基本实体            
            try:
                # 确保可以访问document_dict变量
                document_dict = doc_model.to_dict()
                
                entity = Entity(
                    id=uuid4(),
                    type="concept",
                    name=doc_model.title,
                    description=doc_model.title,
                    properties={},
                    tags=[],
                    importance=None,
                    understanding_level=None,
                    personal_notes=None,
                    category="基础类型",
                    source_id=UUID(doc_model.id),
                    source_type=doc_model.type,
                    source_location=None,
                    extraction_method="manual_upload",
                    confidence=1.0,
                    version=1,
                    previous_version=None,
                    created_at=datetime.now(),
                    updated_at=datetime.now()
                )
                
                # 保存实体到neo4j数据库
                created_entity = await neo4j_db.create(entity)
                print(f"Created document entity: {created_entity.name}")
                
                # 添加到响应中
                return {
                    "document": document_dict,
                    "entity": entity.dict(),
                    "extracted_entities": 0,
                    "extracted_relationships": 0,
                    "message": "Document processed successfully"
                }
            except Exception as e:
                print(f"Error creating document entity: {e}")
                # 如果出错，确保仍然有一个可用的document_dict
                if 'document_dict' not in locals():
                    document_dict = doc_model.to_dict()
                
                # 返回正常响应，不包含实体信息
                return {
                    "document": document_dict,
                    "extracted_entities": 0,
                    "extracted_relationships": 0,
                    "message": "Document processed successfully but entity creation failed"
                }
        # 继续返回原始响应
        # 将文档转为字典以便序列化
        document_dict = doc_model.to_dict()
        
        return {
            "document": document_dict,
            "extracted_entities": len(entities),
            "extracted_relationships": len(relationships),
            "extraction_error": extract_error,
            "message": "Document processed successfully" + (
                " but knowledge extraction failed: " + extract_error if extract_error else ""
            )
        }
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        # Capture any other unexpected errors
        import traceback
        traceback.print_exc()
        print(f"Unexpected error in upload_document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")


@router.post("/url", response_model=dict)
async def process_url(
    background_tasks: BackgroundTasks,
    url: str = Form(...),
    extract_knowledge: bool = Form(False),
    db: Neo4jDatabase = Depends(get_db)
):
    """处理URL并可选地提取知识"""
    try:
        # 创建文档处理器
        document_processor = DocumentProcessor()
        
        # 处理URL
        result = await document_processor.process_url(url)
        
        if result.error:
            raise HTTPException(status_code=400, detail=f"Error processing URL: {result.error}")
        
        # Add to mock documents for testing
        mock_documents.append(result.document)
        
        # Background knowledge extraction similar to document upload
        if extract_knowledge:
            # Implementation similar to upload_document
            pass
        
        return {
            "document": result.document,
            "extracted_entities": 0,
            "extracted_relationships": 0,
            "message": "URL processed successfully"
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing URL: {str(e)}")


@router.get("/", response_model=List[dict])
async def read_documents(
    skip: int = 0,
    limit: int = 100,
    document_type: Optional[str] = None,
    title: Optional[str] = Query(None, description="文档标题（支持模糊匹配）"),
    sqlite_db: Session = Depends(get_sqlite_db)
):
    """获取文档列表，支持分页和过滤"""
    try:
        # 查询SQLite数据库
        query = sqlite_db.query(Document)
        
        # 应用过滤条件
        if document_type:
            query = query.filter(Document.type == document_type)
        
        if title:
            query = query.filter(Document.title.like(f"%{title}%"))
        
        # 按创建时间降序排序
        query = query.order_by(Document.created_at.desc())
        
        # 应用分页
        query = query.offset(skip).limit(limit)
        
        # 获取结果
        documents = query.all()
        print(f"Retrieved {len(documents)} documents from SQLite database")
        
        # 转换为字典列表返回
        result = [doc.to_dict() for doc in documents]
        print(f"Converted {len(result)} documents to dictionary format")
        
        return result
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error retrieving documents: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving documents: {str(e)}")

@router.get("/{document_id}", response_model=dict)
async def read_document(
    document_id: UUID,
    db: Neo4jDatabase = Depends(get_db)
):
    """获取指定ID的文档"""
    # Search in mock documents
    for doc in mock_documents:
        if doc.id == document_id:
            doc_dict = doc.dict()
            doc_dict["id"] = str(doc_dict["id"])
            if "source_id" in doc_dict and doc_dict["source_id"]:
                doc_dict["source_id"] = str(doc_dict["source_id"])
            return doc_dict
    
    raise HTTPException(status_code=404, detail="Document not found")

@router.get("/{document_id}/preview")
async def preview_document(
    document_id: str,
    sqlite_db: Session = Depends(get_sqlite_db)
):
    """预览文档内容"""
    try:
        # 从SQLite获取文档
        document = sqlite_db.query(Document).filter(Document.id == document_id).first()
        
        if document is None:
            raise HTTPException(status_code=404, detail="Document not found")
        
        file_path = document.file_path
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Document file not found")
        
        # 根据文件类型处理预览
        file_type = document.type.lower()
        content = ""
        preview_available = True
        
        # 文本文件处理
        if file_type in ["txt", "text"]:
            try:
                async with aiofiles.open(file_path, 'rb') as f:
                    binary_content = await f.read()
                    
                # 尝试将内容解码为文本
                try:
                    content = binary_content.decode('utf-8', errors='replace')
                except UnicodeDecodeError:
                    content = binary_content.decode('gbk', errors='replace')
            except Exception as e:
                content = f"Error reading file: {str(e)}"
                preview_available = False
        
        # PDF文件处理
        elif file_type == "pdf":
            try:
                # 使用同步方式处理PDF（因为PyPDF2不支持异步）
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text_content = []
                    
                    # 提取前5页或所有页面（如果少于5页）
                    page_limit = min(5, len(pdf_reader.pages))
                    for i in range(page_limit):
                        page = pdf_reader.pages[i]
                        text_content.append(f"==== 第 {i+1} 页 ====\n{page.extract_text()}")
                    
                    # 如果内容太长，进行截断
                    content = "\n\n".join(text_content)
                    if len(content) > 50000:
                        content = content[:50000] + "...\n[内容过长，已截断]"
                    
                    if not content.strip():
                        content = "[PDF 文件] 无法提取文本内容，可能是扫描件或图片PDF。"
            except Exception as e:
                content = f"[PDF 文件] 处理出错: {str(e)}"
                preview_available = False
        
        # Word文档处理
        elif file_type in ["docx", "doc"]:
            if file_type == "docx":
                try:
                    # 使用python-docx处理docx文件
                    import docx
                    doc = docx.Document(file_path)
                    
                    # 提取段落文本
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    
                    # 如果段落太多，只取前100个
                    if len(paragraphs) > 100:
                        paragraphs = paragraphs[:100]
                        paragraphs.append("...\n[内容过长，仅显示前100个段落]")
                    
                    content = "\n\n".join(paragraphs)
                except Exception as e:
                    content = f"[Word 文件] 处理出错: {str(e)}"
                    preview_available = False
            else:
                # .doc格式不支持直接预览
                content = "[Word(.doc) 文件] 无法预览旧版Word格式，请下载后查看。"
                preview_available = False
        
        # 其他类型文件
        else:
            content = f"[{file_type.upper()} 文件] 未知文件类型，无法预览。"
            preview_available = False
        
        return {
            "title": document.title,
            "type": document.type,
            "content": content,
            "preview_available": preview_available
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error previewing document: {str(e)}")

@router.get("/{document_id}/download")
async def download_document(
    document_id: str,
    sqlite_db: Session = Depends(get_sqlite_db)
):
    """下载文档"""
    try:
        # 从SQLite获取文档
        document = sqlite_db.query(Document).filter(Document.id == document_id).first()
        
        if document is None:
            raise HTTPException(status_code=404, detail="Document not found")
        
        file_path = document.file_path
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Document file not found")
        
        # 获取文件名
        filename = os.path.basename(file_path)
        
        # 使用 FileResponse 处理文件下载
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type="application/octet-stream"
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error downloading document: {str(e)}")

@router.delete("/{document_id}", response_model=bool)
async def delete_document(
    document_id: str,
    sqlite_db: Session = Depends(get_sqlite_db)
):
    """删除文档"""
    try:
        # 从SQLite获取文档
        document = sqlite_db.query(Document).filter(Document.id == document_id).first()
        
        if document is None:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # 尝试删除物理文件
        try:
            if document.file_path and os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            if document.archived_path and os.path.exists(document.archived_path):
                os.remove(document.archived_path)
        except Exception as file_error:
            # 记录错误但继续，因为数据库记录已删除
            print(f"Warning: Failed to delete physical file: {str(file_error)}")
        
        # 从数据库删除记录
        sqlite_db.delete(document)
        sqlite_db.commit()
        
        return True
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

# 替换 documents.py 中的 preview_document 函数

@router.get("/{document_id}/preview")
async def preview_document(
    document_id: UUID,
    db: Neo4jDatabase = Depends(get_db)
):
    """预览文档内容"""
    try:
        # 在 mock_documents 中查找
        for doc in mock_documents:
            if doc.id == document_id:
                file_path = doc.file_path
                if not file_path or not os.path.exists(file_path):
                    raise HTTPException(status_code=404, detail="Document file not found")
                
                # 读取文件内容
                content = ""
                try:
                    async with aiofiles.open(file_path, 'rb') as f:
                        binary_content = await f.read()
                        
                    # 尝试将内容解码为文本
                    try:
                        content = binary_content.decode('utf-8', errors='replace')
                    except UnicodeDecodeError:
                        content = f"[Binary content - {len(binary_content)} bytes]"
                        
                except Exception as e:
                    content = f"Error reading file: {str(e)}"
                
                return {
                    "title": doc.title,
                    "type": doc.type,
                    "content": content,
                    "preview_available": True
                }
        
        raise HTTPException(status_code=404, detail="Document not found")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error previewing document: {str(e)}")


# 替换 documents.py 中的 download_document 函数

@router.get("/{document_id}/download")
async def download_document(
    document_id: UUID,
    db: Neo4jDatabase = Depends(get_db)
):
    """下载文档"""
    try:
        # 在 mock_documents 中查找
        for doc in mock_documents:
            if doc.id == document_id:
                file_path = doc.file_path
                if not file_path or not os.path.exists(file_path):
                    raise HTTPException(status_code=404, detail="Document file not found")
                
                filename = os.path.basename(file_path)
                
                # 使用 FileResponse 处理文件下载
                from fastapi.responses import FileResponse
                return FileResponse(
                    path=file_path,
                    filename=filename,
                    media_type="application/octet-stream"
                )
        
        raise HTTPException(status_code=404, detail="Document not found")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error downloading document: {str(e)}")


@router.post("/{document_id}/extract", response_model=Dict[str, Any])
async def extract_knowledge(
    document_id: str,
    sqlite_db: Session = Depends(get_sqlite_db),
    neo4j_db: Neo4jDatabase = Depends(get_db)
):
    """从文档提取知识"""
    try:
        # 获取文档
        document = sqlite_db.query(Document).filter(Document.id == document_id).first()
        
        if document is None:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # 获取文档文件路径
        file_path = document.file_path
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Document file not found")
        
        # 读取文档内容
        content = ""
        try:
            # 读取文件内容（根据文件类型处理）
            file_type = document.type.lower()
            
            # 文本文件
            if file_type in ["txt", "text"]:
                async with aiofiles.open(file_path, 'rb') as f:
                    binary_content = await f.read()
                
                try:
                    content = binary_content.decode('utf-8', errors='replace')
                except UnicodeDecodeError:
                    content = binary_content.decode('gbk', errors='replace')
            
            # PDF文件
            elif file_type == "pdf":
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n\n"
            
            # Word文件
            elif file_type == "docx":
                import docx
                doc = docx.Document(file_path)
                content = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            # 其他文件类型
            else:
                # 尝试作为文本读取
                try:
                    async with aiofiles.open(file_path, 'rb') as f:
                        binary_content = await f.read()
                    content = binary_content.decode('utf-8', errors='replace')
                except:
                    raise HTTPException(
                        status_code=400, 
                        detail=f"Unsupported file type for knowledge extraction: {file_type}"
                    )
        except Exception as e:
            raise HTTPException(
                status_code=500, 
                detail=f"Error reading document content: {str(e)}"
            )
        
        # 创建知识抽取器
        extractor = SpacyNERExtractor(neo4j_db)
        
        # 创建文档对象
        source_document = SourceDocument(
            id=UUID(document.id),
            title=document.title,
            type=document.type,
            content_hash=document.content_hash,
            file_path=document.file_path,
            url=document.url,
            archived_path=document.archived_path,
            metadata=json.loads(document.doc_metadata) if document.doc_metadata else {}
        )
        
        # 提取实体
        entities = await extractor.extract_entities(source_document, content)
        
        # 提取关系
        relationships = []
        if len(entities) >= 2:
            relationships = await extractor.extract_relationships(source_document, entities, content)
        
        # 创建溯源记录
        if entities or relationships:
            await extractor.create_knowledge_traces(source_document, entities, relationships, content)
        
        return {
            "document_id": document_id,
            "extracted_entities": len(entities),
            "extracted_relationships": len(relationships),
            "message": f"成功从文档中提取出 {len(entities)} 个实体和 {len(relationships)} 个关系。"
        }
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"知识提取失败: {str(e)}")


@router.get("/{document_id}/export")
async def export_document(
    document_id: UUID,
    db: Neo4jDatabase = Depends(get_db)
):
    """导出文档及其关联知识"""
    # Find in mock documents
    for doc in mock_documents:
        if doc.id == document_id:
            return {
                "document_id": str(document_id),
                "metadata": doc.metadata,
                "knowledge": {
                    "entities": [],
                    "relationships": []
                }
            }
    
    raise HTTPException(status_code=404, detail="Document not found")